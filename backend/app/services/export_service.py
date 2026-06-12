import io
import csv

from docx import Document
from openpyxl import Workbook
from reportlab.pdfgen import canvas


class ExportService:
    @staticmethod
    def to_csv(records):
        buffer = io.StringIO()
        if not records:
            return "".encode("utf-8")

        fieldnames = list(records[0].keys())
        writer = csv.DictWriter(buffer, fieldnames=fieldnames)
        writer.writeheader()
        writer.writerows(records)
        return buffer.getvalue().encode("utf-8")

    @staticmethod
    def to_excel(records):
        buffer = io.BytesIO()
        workbook = Workbook()
        sheet = workbook.active
        sheet.title = "Report"

        if records:
            headers = list(records[0].keys())
            sheet.append(headers)
            for row in records:
                sheet.append([row.get(header) for header in headers])

        workbook.save(buffer)
        buffer.seek(0)
        return buffer.read()

    @staticmethod
    def to_word(records):
        document = Document()
        document.add_heading("Generated Report", level=1)
        for row in records:
            document.add_paragraph(str(row))
        buffer = io.BytesIO()
        document.save(buffer)
        buffer.seek(0)
        return buffer.read()

    @staticmethod
    def to_pdf(records):
        buffer = io.BytesIO()
        pdf = canvas.Canvas(buffer)
        pdf.drawString(50, 800, "Generated Report")
        y = 780
        for row in records[:30]:
            pdf.drawString(50, y, str(row)[:120])
            y -= 20
            if y < 50:
                pdf.showPage()
                y = 800
        pdf.save()
        buffer.seek(0)
        return buffer.read()
